"""Tests for DOCX file format support in the humanise feature."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest


class TestDocxMissingDependency:
    """Tests for behaviour when python-docx is not installed."""

    def test_missing_dependency_produces_clear_error_message(
        self, tmp_path: Path,
    ) -> None:
        """When python-docx is not installed, humanise_file on a .docx
        should raise ImportError with a message suggesting pip install uwotm8[docx]."""
        docx_file = tmp_path / "test.docx"
        docx_file.write_bytes(b"fake docx content")

        from uwotm8.humanise import humanise_file

        with patch("uwotm8.humanise._has_docx_support", return_value=False):
            with pytest.raises(ImportError, match=r"pip install uwotm8\[docx\]"):
                humanise_file(src=docx_file)


class TestDocxReadWrite:
    """Tests for reading from and writing to .docx files."""

    @pytest.fixture()
    def _require_docx(self) -> None:
        """Skip tests if python-docx is not installed."""
        pytest.importorskip("docx", reason="python-docx not installed")

    @pytest.mark.usefixtures("_require_docx")
    def test_read_text_from_docx_paragraphs_and_runs(self, tmp_path: Path) -> None:
        """Text should be extracted from paragraphs and runs in a .docx file."""
        import docx

        doc = docx.Document()
        doc.add_paragraph("This text has an em dash \u2014 which is an AI tell.")
        doc.add_paragraph("Second paragraph with smart quotes \u201clike this\u201d.")
        docx_path = tmp_path / "input.docx"
        doc.save(str(docx_path))

        from uwotm8.humanise import humanise_file

        result = humanise_file(src=docx_path, check=True)

        # The em dash and smart quotes should be detected as findings.
        tell_names = [f["tell_name"] for f in result.findings]
        assert any("em_dash" in n for n in tell_names), (
            f"Expected em_dash tell in findings, got: {tell_names}"
        )
        assert any("smart" in n and "quote" in n for n in tell_names), (
            f"Expected smart quote tell in findings, got: {tell_names}"
        )

    @pytest.mark.usefixtures("_require_docx")
    def test_write_back_preserves_formatting(self, tmp_path: Path) -> None:
        """After applying tells, bold/italic formatting in the .docx should be preserved."""
        import docx

        doc = docx.Document()
        para = doc.add_paragraph()
        run_bold = para.add_run("Bold text with em dash \u2014 here.")
        run_bold.bold = True
        run_italic = para.add_run(" Italic text\u2019s quote.")
        run_italic.italic = True

        docx_path = tmp_path / "formatted.docx"
        doc.save(str(docx_path))

        from uwotm8.humanise import humanise_file

        result = humanise_file(src=docx_path, check=False)

        # Verify the file was modified (tells were applied).
        assert len(result.findings) > 0

        # Re-open and verify formatting is preserved.
        doc_out = docx.Document(str(docx_path))
        para_out = doc_out.paragraphs[0]
        runs_out = para_out.runs

        # There should still be runs with bold and italic.
        bold_runs = [r for r in runs_out if r.bold]
        italic_runs = [r for r in runs_out if r.italic]
        assert len(bold_runs) > 0, "Bold formatting should be preserved"
        assert len(italic_runs) > 0, "Italic formatting should be preserved"

        # The em dash should have been replaced.
        full_text = para_out.text
        assert "\u2014" not in full_text

    @pytest.mark.usefixtures("_require_docx")
    def test_docx_inline_annotation_not_applicable(self, tmp_path: Path) -> None:
        """DOCX annotation should still produce findings but annotation is
        applied to the text content (not file-level annotation markers)."""
        import docx

        doc = docx.Document()
        doc.add_paragraph("It\u2019s important to note that this has tells.")
        docx_path = tmp_path / "annotate.docx"
        doc.save(str(docx_path))

        from uwotm8.humanise import humanise_file

        # annotate=True should not crash on docx files.
        result = humanise_file(src=docx_path, check=True, annotate=True)
        assert isinstance(result.findings, list)
